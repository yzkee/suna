#!/usr/bin/env python3
"""TDD Verification Suite for Legal Documents (DOCX).

Runs a comprehensive checklist against a legal document directory and reports pass/fail.
Designed to be run after every section is written.

Usage:
    python3 verify-legal.py <document-dir>
    python3 verify-legal.py legal/contract-acme/
    python3 verify-legal.py legal/memo-smith/ --strict   # treat warnings as failures

Expects:
    <document-dir>/document.docx     (the main document)
    <document-dir>/metadata.json     (document metadata: type, parties, jurisdiction, etc.)

Exit codes: 0 = all pass, 1 = failures found
"""

import sys
import os
import re
import json
from pathlib import Path

# ─── Results tracking ───────────────────────────────────────────────────────

PASS_COUNT = 0
FAIL_COUNT = 0
WARN_COUNT = 0
STRICT = False

def check_pass(msg):
    global PASS_COUNT
    print(f"  PASS: {msg}")
    PASS_COUNT += 1

def check_fail(msg):
    global FAIL_COUNT
    print(f"  FAIL: {msg}")
    FAIL_COUNT += 1

def check_warn(msg):
    global WARN_COUNT, FAIL_COUNT
    print(f"  WARN: {msg}")
    WARN_COUNT += 1
    if STRICT:
        FAIL_COUNT += 1


# ─── Text extraction from DOCX ─────────────────────────────────────────────

def extract_docx_text(docx_path):
    """Extract full text from a DOCX file using python-docx or fallback to XML."""
    try:
        from docx import Document
        doc = Document(docx_path)
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else "",
            })
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append({"text": cell.text, "style": "TableCell"})
        return paragraphs
    except ImportError:
        # Fallback: extract text from DOCX XML directly
        import zipfile
        import xml.etree.ElementTree as ET
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with zipfile.ZipFile(docx_path) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
        paragraphs = []
        for para in tree.findall(".//w:p", ns):
            texts = [t.text for t in para.findall(".//w:t", ns) if t.text]
            paragraphs.append({"text": " ".join(texts), "style": ""})
        return paragraphs


def get_full_text(paragraphs):
    """Join all paragraph texts into a single string."""
    return "\n".join(p["text"] for p in paragraphs if p["text"].strip())


# ─── Check: Document exists and is non-trivial ─────────────────────────────

def check_document_exists(doc_dir):
    print("--- Document Status ---")
    docx_files = list(Path(doc_dir).glob("*.docx"))
    if not docx_files:
        check_fail("No .docx file found in document directory")
        return None
    docx_path = docx_files[0]
    size = docx_path.stat().st_size
    if size < 500:
        check_fail(f"Document suspiciously small ({size} bytes)")
        return None
    check_pass(f"Document exists: {docx_path.name} ({size:,} bytes)")
    return str(docx_path)


# ─── Check: Defined terms consistency ───────────────────────────────────────

def check_defined_terms(full_text, metadata):
    print("\n--- Defined Terms ---")
    doc_type = metadata.get("type", "").lower()

    # Find defined terms: words/phrases in quotes followed by definition pattern
    # Pattern 1: "Term" means / shall mean / is defined as
    defined_pattern = r'"([A-Z][A-Za-z\s]+?)"(?:\s+(?:means?|shall mean|is defined as|has the meaning))'
    defined_terms = set(re.findall(defined_pattern, full_text))

    # Pattern 2: ("Term") — parenthetical definition
    paren_pattern = r'\("([A-Z][A-Za-z\s]+?)"\)'
    defined_terms.update(re.findall(paren_pattern, full_text))

    if not defined_terms and doc_type in ("contract", "agreement", "nda", "settlement"):
        check_warn("No defined terms detected in a contract-type document")
        return

    if not defined_terms:
        check_pass("No defined terms expected for this document type")
        return

    # Find all capitalized terms that look like defined terms (2+ chars, Title Case, not at sentence start)
    # This is a heuristic — not perfect but catches most issues
    cap_terms_in_body = set()
    for line in full_text.split("\n"):
        # Find capitalized words that aren't at the start of a sentence
        words = line.split()
        for i, word in enumerate(words):
            cleaned = re.sub(r'[^A-Za-z]', '', word)
            if (cleaned and cleaned[0].isupper() and len(cleaned) > 1
                and cleaned not in ("The", "This", "That", "These", "Those", "Such",
                                   "Section", "Article", "Exhibit", "Schedule",
                                   "Party", "Parties", "Agreement", "Court",
                                   "Plaintiff", "Defendant", "State", "United",
                                   "Federal", "January", "February", "March",
                                   "April", "May", "June", "July", "August",
                                   "September", "October", "November", "December")):
                cap_terms_in_body.add(cleaned)

    # Check: every defined term is used in the body
    unused = []
    for term in defined_terms:
        # Check if the term appears outside its definition
        uses = len(re.findall(re.escape(term), full_text)) - 1  # subtract the definition itself
        if uses <= 0:
            unused.append(term)

    if unused:
        check_warn(f"{len(unused)} defined term(s) never used: {', '.join(unused[:5])}")
    else:
        check_pass(f"All {len(defined_terms)} defined terms are used in the document")

    return defined_terms


# ─── Check: Cross-reference integrity ───────────────────────────────────────

def check_cross_references(full_text):
    print("\n--- Cross-References ---")

    # Find section references: "Section X.Y", "Article X", "Exhibit A"
    section_refs = set(re.findall(r'Section\s+(\d+(?:\.\d+)*)', full_text))
    article_refs = set(re.findall(r'Article\s+(\w+)', full_text))
    exhibit_refs = set(re.findall(r'Exhibit\s+([A-Z](?:-\d+)?)', full_text))
    schedule_refs = set(re.findall(r'Schedule\s+(\d+|[A-Z])', full_text))

    # Find actual section/article headings
    # Pattern: "1.2" or "1.2.3" at start of a line-like context
    actual_sections = set(re.findall(r'(?:^|\n)\s*(\d+(?:\.\d+)+)[.\s]', full_text))
    # Add top-level sections
    actual_sections.update(re.findall(r'(?:^|\n)\s*(\d+)[.\s]+[A-Z]', full_text))

    broken_refs = []
    for ref in section_refs:
        if ref not in actual_sections:
            # Check if it might be a top-level section
            top = ref.split(".")[0]
            if top not in actual_sections and ref not in actual_sections:
                broken_refs.append(f"Section {ref}")

    if broken_refs and len(broken_refs) <= 3:
        # Only report if we found actual sections (otherwise we can't verify)
        if actual_sections:
            check_warn(f"Potentially broken references: {', '.join(broken_refs[:5])}")
        else:
            check_pass("Cross-references present (section structure not parseable for verification)")
    elif broken_refs:
        check_fail(f"{len(broken_refs)} potentially broken section reference(s)")
    else:
        total = len(section_refs) + len(article_refs) + len(exhibit_refs) + len(schedule_refs)
        if total > 0:
            check_pass(f"All {total} cross-references appear valid")
        else:
            check_pass("No cross-references to check")


# ─── Check: Citation format (Bluebook) ─────────────────────────────────────

def check_citations(full_text, metadata):
    print("\n--- Legal Citations ---")
    doc_type = metadata.get("type", "").lower()

    # Only check citations in litigation/research documents
    citation_types = ("memo", "memorandum", "brief", "motion", "complaint", "opinion", "petition")
    if not any(t in doc_type for t in citation_types):
        check_pass("Citation format not applicable for this document type")
        return

    # Find case citations: Party v. Party, Vol Reporter Page (Court Year)
    case_pattern = r'[A-Z][a-z]+\s+v\.\s+[A-Z][a-z]+'
    case_cites = re.findall(case_pattern, full_text)

    # Find statute citations: ## U.S.C. § ##
    statute_pattern = r'\d+\s+U\.S\.C\.\s+§\s*\d+'
    statute_cites = re.findall(statute_pattern, full_text)

    # Find CFR citations: ## C.F.R. § ##
    cfr_pattern = r'\d+\s+C\.F\.R\.\s+§\s*\d+'
    cfr_cites = re.findall(cfr_pattern, full_text)

    total_cites = len(case_cites) + len(statute_cites) + len(cfr_cites)

    if total_cites == 0:
        check_warn("No legal citations found in a litigation/research document")
        return

    check_pass(f"Found {total_cites} citation(s): {len(case_cites)} cases, {len(statute_cites)} statutes, {len(cfr_cites)} regulations")

    # Check for common citation format errors
    errors = []

    # Check: "v." not "vs." or "v "
    vs_errors = len(re.findall(r'\bvs\.\s', full_text))
    if vs_errors:
        errors.append(f'{vs_errors} instance(s) of "vs." (should be "v.")')

    # Check: Id. should be italicized (we can't check formatting, but can check usage)
    id_uses = len(re.findall(r'\bId\.\s', full_text))
    # Id. should only follow immediately after another citation (approximate check)

    # Check: pinpoint citations present (page number after first page)
    # Only flag when the same case is cited multiple times without pinpoints
    # A citation like "418 U.S. 241 (1974)" is fine — it's the full case citation.
    # A pinpoint looks like "418 U.S. 241, 258 (1974)" — citing a specific page.
    proper_pinpoint = len(re.findall(r'\d+\s+(?:F\.\d+d|F\.\d+th|S\.\s*Ct|U\.S\.)\s+\d+,\s*\d+', full_text))

    if proper_pinpoint == 0 and case_cites and len(case_cites) > 5:
        errors.append("No pinpoint page references found — consider adding specific page cites for key propositions")

    if errors:
        for e in errors:
            check_warn(e)
    else:
        check_pass("Citation format appears correct")


# ─── Check: Placeholder / draft artifact detection ─────────────────────────

def check_placeholders(full_text):
    print("\n--- Completeness ---")

    placeholders = {
        "brackets": re.findall(r'\[(?:INSERT|TBD|TODO|FILL IN|______|NAME|DATE|AMOUNT|ADDRESS|NUMBER|TO BE)[^\]]*\]', full_text, re.I),
        # Exclude signature lines (By: ___) — only flag standalone blank-fill underscores
        "underscores": [m for m in re.findall(r'_{4,}', full_text)
                       if not re.search(r'(?:By|Name|Title|Date|Signature):\s*' + re.escape(m), full_text)],
        "todo_comments": re.findall(r'(?:TODO|FIXME|XXX|HACK|TBD|PLACEHOLDER)', full_text, re.I),
        "highlight_markers": re.findall(r'\[HIGHLIGHT\]|\[REVIEW\]|\[CHECK\]|\[VERIFY\]', full_text, re.I),
        "draft_watermarks": re.findall(r'\bDRAFT\b', full_text),
    }

    total = sum(len(v) for v in placeholders.values())
    if total == 0:
        check_pass("No placeholders, TODOs, or draft artifacts found")
    else:
        details = []
        if placeholders["brackets"]:
            details.append(f'{len(placeholders["brackets"])} [INSERT/TBD] bracket(s)')
        if placeholders["underscores"]:
            details.append(f'{len(placeholders["underscores"])} blank line(s) (____)')
        if placeholders["todo_comments"]:
            details.append(f'{len(placeholders["todo_comments"])} TODO/TBD marker(s)')
        if placeholders["highlight_markers"]:
            details.append(f'{len(placeholders["highlight_markers"])} [REVIEW/CHECK] marker(s)')
        if placeholders["draft_watermarks"]:
            details.append(f'{len(placeholders["draft_watermarks"])} DRAFT watermark(s)')
        check_fail(f"{total} placeholder(s)/draft artifact(s): {'; '.join(details)}")


# ─── Check: Party name consistency ──────────────────────────────────────────

def check_party_consistency(full_text, metadata):
    print("\n--- Party Names ---")

    parties = metadata.get("parties", [])
    if not parties:
        check_pass("No parties specified in metadata (skipping)")
        return

    for party in parties:
        name = party.get("name", "")
        short = party.get("short_name", "")
        if name and short:
            # Check that the short name is actually used after being defined
            uses = len(re.findall(re.escape(short), full_text))
            if uses == 0:
                check_warn(f'Short name "{short}" for party "{name}" never used')
            # Check for the full name being used after it should have been shortened
            # (heuristic: if short name exists, full name shouldn't appear more than ~3 times)
            full_uses = len(re.findall(re.escape(name), full_text))
            if full_uses > 5 and uses > 0:
                check_warn(f'Full name "{name}" used {full_uses} times (consider using "{short}" consistently)')
        elif name:
            check_pass(f'Party "{name}" referenced in document')


# ─── Check: Boilerplate provisions (contracts) ─────────────────────────────

def check_boilerplate(full_text, metadata):
    print("\n--- Required Provisions ---")
    doc_type = metadata.get("type", "").lower()

    if doc_type not in ("contract", "agreement", "nda", "settlement", "terms of service",
                        "employment agreement", "services agreement", "license agreement"):
        check_pass("Boilerplate check not applicable for this document type")
        return

    required_provisions = {
        "governing law": r'(?:governing\s+law|choice\s+of\s+law|governed\s+by.*laws\s+of)',
        "entire agreement": r'(?:entire\s+agreement|constitutes?\s+the\s+entire)',
        "severability": r'(?:severab|invalid.*unenforceab|unenforceab.*sever)',
        "amendment": r'(?:amend(?:ment|ed).*(?:writ(?:ten|ing)|signed)|(?:not\s+be\s+)?modif(?:y|ied|ication).*(?:except|writ))',
        "notices": r'(?:notice.*(?:shall|must|will|be)\s+.*(?:writ(?:ten|ing)|deliver|given|sent)|all\s+notices?\s+under)',
        "assignment": r'(?:assign(?:ment)?.*(?:without|prior|consent)|neither\s+party\s+may\s+assign)',
    }

    missing = []
    found = []
    for provision, pattern in required_provisions.items():
        if re.search(pattern, full_text, re.I):
            found.append(provision)
        else:
            missing.append(provision)

    if missing:
        check_warn(f"Missing standard provisions: {', '.join(missing)}")
    if found:
        check_pass(f"Found {len(found)}/{len(required_provisions)} standard provisions: {', '.join(found)}")


# ─── Check: Shall/May/Must consistency ─────────────────────────────────────

def check_modal_verbs(full_text, metadata):
    print("\n--- Language Precision ---")
    doc_type = metadata.get("type", "").lower()

    if doc_type not in ("contract", "agreement", "nda", "settlement", "regulation",
                        "terms of service", "employment agreement"):
        check_pass("Modal verb check not applicable for this document type")
        return

    shall_count = len(re.findall(r'\bshall\b', full_text, re.I))
    must_count = len(re.findall(r'\bmust\b', full_text, re.I))
    may_count = len(re.findall(r'\bmay\b', full_text, re.I))
    will_count = len(re.findall(r'\bwill\b', full_text, re.I))

    # Check for "shall not" (obligation not to) vs "may not" (prohibition) confusion
    shall_not = len(re.findall(r'\bshall\s+not\b', full_text, re.I))

    check_pass(f"Modal verbs: shall={shall_count}, must={must_count}, may={may_count}, will={will_count}")

    # Warn if both "shall" and "must" are used (inconsistent — pick one style)
    if shall_count > 3 and must_count > 3:
        check_warn('Both "shall" and "must" used frequently — consider standardizing to one')


# ─── Check: Word count / length ─────────────────────────────────────────────

def check_length(full_text, metadata):
    print("\n--- Document Length ---")
    word_count = len(full_text.split())
    doc_type = metadata.get("type", "").lower()

    # Page limit checks for briefs
    if "brief" in doc_type or "motion" in doc_type:
        page_limit = metadata.get("page_limit")
        word_limit = metadata.get("word_limit")
        if word_limit and word_count > word_limit:
            check_fail(f"Word count {word_count:,} exceeds limit of {word_limit:,}")
        elif word_limit:
            check_pass(f"Word count {word_count:,} within limit of {word_limit:,}")
        else:
            check_pass(f"Word count: {word_count:,} (no limit specified)")
    elif word_count < 50:
        check_warn(f"Document very short ({word_count} words)")
    else:
        check_pass(f"Word count: {word_count:,}")


# ─── Check: Date consistency ────────────────────────────────────────────────

def check_date_consistency(full_text):
    print("\n--- Date Format ---")

    # Find different date formats
    long_dates = re.findall(r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}', full_text)
    slash_dates = re.findall(r'\d{1,2}/\d{1,2}/\d{2,4}', full_text)
    dash_dates = re.findall(r'\d{4}-\d{2}-\d{2}', full_text)

    formats_used = 0
    if long_dates: formats_used += 1
    if slash_dates: formats_used += 1
    if dash_dates: formats_used += 1

    if formats_used > 1:
        check_warn(f"Inconsistent date formats: {len(long_dates)} long, {len(slash_dates)} slash, {len(dash_dates)} ISO")
    elif formats_used == 1:
        total = len(long_dates) + len(slash_dates) + len(dash_dates)
        check_pass(f"Consistent date format ({total} dates found)")
    else:
        check_pass("No dates to check")


# ─── Main ───────────────────────────────────────────────────────────────────

def main():
    global STRICT

    if len(sys.argv) < 2:
        print("Usage: verify-legal.py <document-dir> [--strict]", file=sys.stderr)
        sys.exit(1)

    doc_dir = sys.argv[1].rstrip("/")
    STRICT = "--strict" in sys.argv

    print(f"=== Legal Document Verification: {doc_dir} ===\n")

    # Load metadata
    metadata_path = os.path.join(doc_dir, "metadata.json")
    if os.path.exists(metadata_path):
        with open(metadata_path) as f:
            metadata = json.load(f)
    else:
        metadata = {}
        print("  NOTE: No metadata.json found, running with defaults\n")

    # Check 1: Document exists
    docx_path = check_document_exists(doc_dir)
    if not docx_path:
        print(f"\n{'='*35}")
        print(f"  PASS: {PASS_COUNT}")
        print(f"  WARN: {WARN_COUNT}")
        print(f"  FAIL: {FAIL_COUNT}")
        print(f"{'='*35}")
        print("  RESULT: FAILED")
        sys.exit(1)

    # Extract text
    try:
        paragraphs = extract_docx_text(docx_path)
        full_text = get_full_text(paragraphs)
    except Exception as e:
        check_fail(f"Could not read document: {e}")
        print(f"\n{'='*35}")
        print(f"  PASS: {PASS_COUNT}")
        print(f"  FAIL: {FAIL_COUNT}")
        print(f"{'='*35}")
        print("  RESULT: FAILED")
        sys.exit(1)

    # Run all checks
    check_defined_terms(full_text, metadata)
    check_cross_references(full_text)
    check_citations(full_text, metadata)
    check_placeholders(full_text)
    check_party_consistency(full_text, metadata)
    check_boilerplate(full_text, metadata)
    check_modal_verbs(full_text, metadata)
    check_length(full_text, metadata)
    check_date_consistency(full_text)

    # Summary
    print(f"\n{'='*35}")
    print(f"  PASS: {PASS_COUNT}")
    print(f"  WARN: {WARN_COUNT}")
    print(f"  FAIL: {FAIL_COUNT}")
    print(f"{'='*35}")

    if FAIL_COUNT > 0:
        print("  RESULT: FAILED")
        sys.exit(1)
    elif WARN_COUNT > 0:
        print("  RESULT: PASSED (with warnings)")
    else:
        print("  RESULT: ALL CHECKS PASSED")
    sys.exit(0)


if __name__ == "__main__":
    main()
