"""
Export API for converting HTML/Markdown to PDF and DOCX

This module provides endpoints for exporting content from the mobile/web app
to various formats (PDF, DOCX, HTML, Markdown).
"""

from pathlib import Path
from io import BytesIO
from typing import Optional

from fastapi import APIRouter, HTTPException, Depends, Request
from fastapi.responses import Response
from pydantic import BaseModel, Field
from urllib.parse import quote
from core.utils.auth_utils import verify_and_get_user_id_from_jwt

try:
    from weasyprint import HTML, CSS
except (ImportError, OSError) as e:
    weasyprint_available = False
    print(f"[WARNING] WeasyPrint not available: {e}")
    print("[INFO] To fix on macOS, run: export DYLD_FALLBACK_LIBRARY_PATH=$(brew --prefix)/lib:$DYLD_FALLBACK_LIBRARY_PATH")
else:
    weasyprint_available = True

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx_available = False
else:
    docx_available = True

try:
    from bs4 import BeautifulSoup
except ImportError:
    beautifulsoup_available = False
else:
    beautifulsoup_available = True


router = APIRouter(prefix="/export", tags=["export"])


class ExportRequest(BaseModel):
    """Request model for export endpoints"""
    content: str = Field(..., description="HTML content to export")
    fileName: str = Field(..., description="Base filename (without extension)")


class ExportResponse(BaseModel):
    """Response model for export endpoints"""
    success: bool
    message: str
    filename: Optional[str] = None


# Professional PDF styles matching the Next.js implementation
PDF_STYLES = """
  * { 
    margin: 0; 
    padding: 0; 
    box-sizing: border-box; 
  }
  
  body { 
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #1a1a1a;
    padding: 0;
    margin: 0;
    background: #fff;
  }

  h1, h2, h3, h4, h5, h6 {
    font-weight: 600;
    line-height: 1.3;
    color: #111;
    letter-spacing: -0.02em;
    page-break-after: avoid;
  }
  
  h1 { 
    font-size: 18pt; 
    margin: 24pt 0 12pt 0;
    padding-bottom: 6pt;
    border-bottom: 1pt solid #e0e0e0;
  }
  h1:first-child { margin-top: 0; }
  
  h2 { font-size: 15pt; margin: 20pt 0 10pt 0; }
  h2:first-child { margin-top: 0; }
  
  h3 { font-size: 13pt; margin: 16pt 0 8pt 0; }
  h3:first-child { margin-top: 0; }
  
  h4 { font-size: 11pt; margin: 14pt 0 6pt 0; }
  h4:first-child { margin-top: 0; }
  
  h5 { font-size: 11pt; margin: 12pt 0 4pt 0; }
  h5:first-child { margin-top: 0; }
  
  h6 { 
    font-size: 10pt; 
    margin: 12pt 0 4pt 0;
    color: #555;
    text-transform: uppercase;
    letter-spacing: 0.5pt;
  }
  h6:first-child { margin-top: 0; }

  p { margin: 8pt 0; line-height: 1.65; }
  p:first-child { margin-top: 0; }
  p:last-child { margin-bottom: 0; }
  
  strong, b { font-weight: 600; }
  em, i { font-style: italic; }
  u { text-decoration: underline; }
  s, strike, del { text-decoration: line-through; color: #666; }

  a { color: #0066cc; text-decoration: underline; }

  ul, ol { margin: 10pt 0; padding-left: 28pt; }
  ul:first-child, ol:first-child { margin-top: 0; }
  ul:last-child, ol:last-child { margin-bottom: 0; }
  
  li { margin: 4pt 0; line-height: 1.5; padding-left: 4pt; }
  li > ul, li > ol { margin: 4pt 0; }

  blockquote { 
    margin: 12pt 0;
    padding: 8pt 12pt;
    border-left: 3pt solid #e0e0e0;
    background: #fafafa;
    color: #555;
    font-style: italic;
    border-radius: 0 4pt 4pt 0;
  }

  code {
    font-family: 'Courier New', Consolas, 'Lucida Console', monospace;
    font-size: 10pt;
    background: #f4f4f5;
    padding: 2pt 4pt;
    border-radius: 2pt;
    color: #1a1a1a;
  }
  
  pre {
    margin: 12pt 0;
    padding: 12pt;
    background: #f4f4f5;
    border: 1pt solid #e4e4e7;
    border-radius: 4pt;
    overflow-x: auto;
    font-size: 9.5pt;
    line-height: 1.5;
    page-break-inside: avoid;
  }
  pre code {
    background: none;
    padding: 0;
    border-radius: 0;
  }

  table { 
    width: 100%;
    margin: 14pt 0;
    border-collapse: collapse; 
    font-size: 10.5pt;
    page-break-inside: avoid;
  }
  
  th { 
    padding: 8pt 10pt;
    border: 1pt solid #d0d0d0;
    text-align: left;
    font-weight: 600;
    background: #f8f9fa;
  }
  
  td { 
    padding: 8pt 10pt;
    border: 1pt solid #d0d0d0;
    text-align: left;
    vertical-align: top;
  }

  hr { border: none; border-top: 1pt solid #e0e0e0; margin: 20pt 0; }

  img { max-width: 100%; height: auto; display: block; margin: 12pt 0; page-break-inside: avoid; }

  @media print {
    body { font-size: 11pt; -webkit-print-color-adjust: exact; print-color-adjust: exact; }
    pre { white-space: pre-wrap; word-wrap: break-word; page-break-inside: avoid; }
    h1, h2, h3, h4, h5, h6 { page-break-after: avoid; }
    table { page-break-inside: avoid; }
    img { page-break-inside: avoid; }
    tr { page-break-inside: avoid; }
  }

  @page { size: A4; margin: 20mm; }
"""


def sanitize_filename(name: str) -> str:
    """Sanitize filename for safe file system use"""
    import re
    return re.sub(r'[^\w\s-]', '', name).strip()[:100] or 'document'


def preprocess_html(html: str) -> str:
    """Preprocess HTML for better conversion"""
    import re
    
    if not html or not isinstance(html, str):
        return '<p></p>'
    
    # Remove script/style tags
    html = re.sub(r'<script\b[^<]*(?:(?!<\/script>)<[^<]*)*<\/script>', '', html, flags=re.IGNORECASE)
    html = re.sub(r'<style\b[^<]*(?:(?!<\/style>)<[^<]*)*<\/style>', '', html, flags=re.IGNORECASE)
    
    # Clean up empty paragraphs
    html = re.sub(r'<p>\s*<\/p>', '', html)
    
    return html.strip()


@router.post("/pdf")
async def export_to_pdf(
    request: ExportRequest,
    user_id: str = Depends(verify_and_get_user_id_from_jwt)
):
    """
    Export HTML content to PDF using WeasyPrint
    
    Requires authentication.
    Returns the PDF file directly for download.
    """
    if not weasyprint_available:
        raise HTTPException(
            status_code=503,
            detail="PDF export is not available. WeasyPrint is not installed."
        )
    
    try:
        content = request.content
        file_name = sanitize_filename(request.fileName)
        
        print(f"[PDF Export] User: {user_id}, File: {file_name}, Content length: {len(content)}")
        
        # Preprocess HTML
        preprocessed_html = preprocess_html(content)
        
        # Create full HTML document with styles
        full_html = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>{file_name}</title>
  <style>{PDF_STYLES}</style>
</head>
<body>
{preprocessed_html}
</body>
</html>"""
        
        # Generate PDF with WeasyPrint
        pdf_bytes = HTML(string=full_html).write_pdf()
        
        # Return PDF file
        encoded_filename = quote(f"{file_name}.pdf", safe="")
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )
                
    except Exception as e:
        print(f"❌ PDF export error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate PDF: {str(e)}"
        )


@router.post("/docx")
async def export_to_docx(
    request: ExportRequest,
    user_id: str = Depends(verify_and_get_user_id_from_jwt)
):
    """
    Export HTML content to DOCX format
    
    Requires authentication.
    Returns the DOCX file directly for download.
    """
    if not docx_available or not beautifulsoup_available:
        raise HTTPException(
            status_code=503,
            detail="DOCX export is not available. Required libraries are not installed."
        )
    
    try:
        content = request.content
        file_name = sanitize_filename(request.fileName)
        
        print(f"[DOCX Export] User: {user_id}, File: {file_name}, Content length: {len(content)}")
        
        # Create DOCX document with clean black/white styling
        doc = Document()
        
        # Set default font to clean black/white theme
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(26, 26, 26)  # #1a1a1a - dark gray/black
        
        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')
        
        # Convert HTML elements to DOCX
        for element in soup.children:
            process_html_element(element, doc)
        
        # Save to BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return DOCX file
        encoded_filename = quote(f"{file_name}.docx", safe="")
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )
        
    except Exception as e:
        print(f"❌ DOCX export error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate DOCX: {str(e)}"
        )


def process_html_element(element, doc, parent_paragraph=None):
    """Process HTML element and add to DOCX document"""
    if element.name is None:
        text = str(element).strip()
        if text and parent_paragraph:
            parent_paragraph.add_run(text)
        return

    if element.name == 'p':
        p = doc.add_paragraph()
        for child in element.children:
            process_inline_element(child, p)
            
    elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1])
        heading = doc.add_heading(element.get_text(), level)
        # Ensure heading is black
        for run in heading.runs:
            run.font.color.rgb = RGBColor(17, 17, 17)  # #111 - dark black for headings
        
    elif element.name == 'ul':
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            for child in li.children:
                process_inline_element(child, p)
                
    elif element.name == 'ol':
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Number')
            for child in li.children:
                process_inline_element(child, p)
                
    elif element.name == 'blockquote':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(element.get_text())
        run.font.italic = True
        run.font.color.rgb = RGBColor(85, 85, 85)  # #555 - gray for blockquotes
        
    elif element.name == 'pre':
        p = doc.add_paragraph()
        run = p.add_run(element.get_text())
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(26, 26, 26)  # Black for code blocks
        
    elif element.name == 'table':
        process_table_element(element, doc)


def process_inline_element(element, paragraph):
    """Process inline HTML element and add to paragraph"""
    if element.name is None:
        text = str(element).strip()
        if text:
            paragraph.add_run(text)
        return
    
    if element.name == 'strong' or element.name == 'b':
        run = paragraph.add_run(element.get_text())
        run.bold = True
        
    elif element.name == 'em' or element.name == 'i':
        run = paragraph.add_run(element.get_text())
        run.italic = True
        
    elif element.name == 'u':
        run = paragraph.add_run(element.get_text())
        run.underline = True
        
    elif element.name == 'code':
        run = paragraph.add_run(element.get_text())
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(26, 26, 26)  # Black for inline code
        
    elif element.name == 'a':
        run = paragraph.add_run(element.get_text())
        run.font.color.rgb = RGBColor(26, 26, 26)  # Black instead of blue
        run.underline = True
        
    else:
        for child in element.children:
            process_inline_element(child, paragraph)


def process_table_element(table_element, doc):
    """Process HTML table and add to DOCX document"""
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    if max_cols == 0:
        return
    
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                table_cell = table.rows[i].cells[j]
                table_cell.text = cell.get_text().strip()
                
                if cell.name == 'th':
                    for paragraph in table_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(17, 17, 17)  # Black for table headers


@router.get("/health")
async def export_health_check():
    """Export service health check endpoint"""
    return {
        "status": "healthy",
        "service": "export-api",
        "pdf_available": weasyprint_available,
        "docx_available": docx_available and beautifulsoup_available,
    }
